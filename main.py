import os
import sys
import logging
import time
from datetime import datetime
from typing import Dict, List

import docx
import pandas as pd
from docxcompose.composer import Composer

from config import (
    CONTRACT_DOC,
    EXPERT_LIST,
    RECEIPT_DOC,
    SIGNATURE_DOC,
    TAIWAN_DATE_OFFSET,
    TARGET_DIR,
)
from fs_monitor import init_fs_handler


def init_logger(filename: str = "logs/logs.log") -> logging.RootLogger:
    logger = logging.getLogger(__name__)
    logging.basicConfig(filename=filename, encoding="utf-8", level=logging.DEBUG)
    handler = logging.StreamHandler(sys.stdout)
    logger.addHandler(handler)
    return logger


def merge_col_to_string(col: pd.Series, sep="_") -> str:
    """
    Merge pandas series values into single string.
    Args:
        col (pd.Series): Column values to return as single string
    Returns:
        Str containing values (e.g. Series[1, 2, 3] -> "1_2_3")
    """
    return sep.join(col.values)


def compose_save(docs: List[docx.Document], filename: str):
    """
    Combine the documents into a single word file and save.
    Args:
        docs (List[docx.Document]): List of edited Word documents.
        filename (str): Location to save the merged doc.
    """
    assert len(docs) >= 1, "Empty document list, make sure at least 1!"
    composer = None
    for doc in docs:
        if composer is None:
            composer = Composer(doc)
            master = doc
        else:
            master.add_page_break()
            composer.append(doc)
    composer.save(filename)


def format_contract_fill_in_data(data: pd.DataFrame) -> List[Dict[str, str]]:
    """
    Organize the dict of values we search for and the data we fill in.
    Args:
        data (pd.DataFrame): User data
    Returns:
        List of dictionaries with values to search for and fill in for each user.
    """
    formatted = []
    conv_date = convert_date_to_chinese(data.iloc[0]["會議日期"].strftime("%Y-%m-%d"))
    for index, row in data.iterrows():
        formatted_row = {
            "意於年月日": f"意於{conv_date}",
            "申請案號：案號": f"申請案號：{row['案號']}",
            "課程名稱：課程全名": f"課程名稱：{row['課程名稱']}",
            "單位名稱：單位全名": f"單位名稱：{row['單位名稱']}",
            "立切結書人：姓名": f"立切結書人：{row['姓名']}",
            "身分證統一編號：身分證字號": f"身分證統一編號：{row['身分證字號']}",
            "中華民國Date": f"中華民國{conv_date}",
        }
        formatted.append(formatted_row)
    return formatted


def format_receipt_fill_in_data(data: pd.DataFrame) -> List[Dict[str, str]]:
    """
    Organize the dictionary values to search for and the values to replace for receipt file.
    Args:
        data (pd.DataFrame): User data
    Returns:
        List of dictionaries with values to search for and fill in for each user.
    """
    formatted = []
    for index, row in data.iterrows():
        date = convert_date_to_chinese(row["會議日期"].strftime("%Y-%m-%d"))
        formatted_row = {
            "姓名OOO": f"{row['姓名']}",
            "電話OOOOOOOOOO": f"{row['手機']}",
            "Date": f"{date}",
            "住址：OOOOOOOOOOO": f"住址:{row['郵遞區號-通訊地址']}",
            "中華民國國籍：身分證統一編號　IDOOOOOOOOOO": f"■ 中華民國國籍：身分證統一編號{row['身分證字號']}",
        }
        formatted.append(formatted_row)
    return formatted


def convert_date_to_chinese(date: str, separator: str = None) -> str:
    """
    Take a DateTime object and format it to Taiwanese format.
    Args:
        date (datetime.TimeStamp): DateTime object with needed date
        separator (str): Character separating the items in a date, e.g. "113.12.25", default "年月日"
    Returns:
        (str) containing the time converted to Taiwanese time (e.g 1998/03/29 -> 87年03月29日)
    """
    if separator:
        year = month = separator
        day = ""
    else:
        year, month, day = "年", "月", "日"

    dates = date.split("-")
    assert len(dates) == 3, "Not enough values in the date."
    roc_year = int(dates[0]) - TAIWAN_DATE_OFFSET
    # Remove **only** leading zeros from months, days
    dates[1] = dates[1][1:] if dates[1][0] == "0" else dates[1]
    dates[2] = dates[2][1:] if dates[2][0] == "0" else dates[2]
    return f"{roc_year}{year}{dates[1]}{month}{dates[2]}{day}"


def search_and_replace_expert_info(
    filename: str,
    user_data: List[Dict[str, str]],
) -> List[docx.Document]:
    """
    Search for keywords in the document's paragraphs and replace it with the matching data.
    Args:
        filename (str): Name of the Word file to open
        user_data (List[Dict[str, str]]): Formatted information for each user, phrase
            to search for and value to replace by.
    Returns:
        List of documents containing the modified word documents.
    """
    modified_docs = []
    for idx, user in enumerate(user_data):
        doc = docx.Document(filename)
        for par in doc.paragraphs:
            for keyword, value in user.items():
                if keyword in par.text:
                    to_replace = par.text.replace(keyword, value)
                    for run in par.runs:
                        run.text = ""
                    par.runs[0].text = to_replace

        modified_docs.append(doc)
    return modified_docs


def format_signature_sheet_fill_in_data(data: pd.DataFrame) -> Dict[str, str]:
    """
    Return the string with the name and date to fill in, only need one copy for the entire group.
    """
    # Use data only from first expert, necessary fields should be the same for all
    row = data.iloc[0]
    date = convert_date_to_chinese(row["會議日期"].strftime("%Y-%m-%d"))
    day_of_week = row["會議日期"].weekday()
    weekday_mapping = {0: "一", 1: "二", 2: "三", 3: "四", 4: "五", 5: "六", 6: "日"}
    hour = row["會議時間(24時)"].hour
    minute = row["會議時間(24時)"].minute

    time_of_day = "上"
    if hour > 12:
        hour -= 12
        time_of_day = "下"
    elif hour == 12:  # At 12, it's afternoon but we don't subtract 12
        time_of_day = "下"

    # Only print the minutes if nonzero
    if minute > 0:
        minute_of_day = f"{minute}分鐘"
    else:
        minute_of_day = ""

    formatted = {
        "貳、時間：Date": f"貳、時間：{date} ({weekday_mapping[day_of_week]}) {time_of_day}午{hour}時{minute_of_day}",
        "肆、審查案件：Number": f"肆、審查案件：{row['案號'] + " " + row['課程名稱']}",
    }
    return formatted


def edit_contract(data: pd.DataFrame):
    """
    Edit 切結書, fill in date, number, and expert information
    Args:
        data (pd.DataFrame): dataframe from which to fill columns.
    """

    formatted = format_contract_fill_in_data(data)
    docs = search_and_replace_expert_info(CONTRACT_DOC, formatted)

    compose_save(docs, os.path.join(TARGET_DIR, f"{data.iloc[0]['案號']}_切結書.docx"))


def edit_receipt(data: pd.DataFrame):
    """
    Edit the receipt file  領據 with expert information.
    Args:
        data (pd.DataFrame): DataFrame containing expert information
    """
    formatted = format_receipt_fill_in_data(data)
    docs = search_and_replace_expert_info(RECEIPT_DOC, formatted)

    compose_save(docs, os.path.join(TARGET_DIR, f"{data.iloc[0]['案號']}_領據.docx"))


def edit_signature_sheet(data: pd.DataFrame):
    """
    Edit the signature sheet with the data for all the teachers of a course.
    Args:
        data (pd.DataFrame): DataFrame with all expert's information
    """
    formatted = format_signature_sheet_fill_in_data(data)
    doc = search_and_replace_expert_info(SIGNATURE_DOC, [formatted])
    table = doc[0].tables[0]
    for idx, expert in data.iterrows():
        row = table.rows[idx + 1].cells  # Offset header row
        row[1].paragraphs[0].runs[0].text = expert["現職單位"]
        row[2].paragraphs[0].runs[0].text = expert["職稱"]  # Center
        row[3].paragraphs[0].runs[0].text = expert["姓名"]

    # Print date on the bottom table
    date = data.iloc[0]["會議日期"].strftime("%Y-%m-%d")
    conv_date = convert_date_to_chinese(date, separator=".")
    table = doc[0].tables[1]
    table.rows[0].cells[3].paragraphs[0].runs[0].text = conv_date

    case = data.iloc[0]["案號"]
    doc[0].save(os.path.join(TARGET_DIR, f"{case}_簽到表.docx"))


def main(**kwargs):
    """
    Initial function called on startup.
    Process the Excel document and fill in the fields on
    the required word docs.
    """
    start = time.time()
    logger = init_logger()
    logger.info(f"Starting program at {datetime.now()}")

    # TODO: Change to support other Excel names
    expert_file = EXPERT_LIST if "EXPERT_LIST" not in kwargs else kwargs["EXPERT_LIST"]

    try:
        expert_info = pd.read_excel(expert_file)
        logger.info(f"Loaded {expert_file}")
    except FileNotFoundError or PermissionError as e:
        logger.error("Problem opening file: {e}")
        return

    try:
        edit_contract(expert_info)
        edit_receipt(expert_info)
        edit_signature_sheet(expert_info)
        end = time.time()
        logger.info(f"Finsihed in {end - start} sec.")
    except Exception as e:
        logger.error(f"Error running main: {e}")


if __name__ == "__main__":
    run_pipeline = False
    if run_pipeline:
        main()
    else:
        init_fs_handler(main)
